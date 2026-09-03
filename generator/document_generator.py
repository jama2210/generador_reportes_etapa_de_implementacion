from io import BytesIO

from docx import Document
from docx.shared import Inches

from config import (
    COL_FECHA,
    COL_NOMBRE_ASESORIA,
    LOGO_PATH,
    MODALITY_SECTIONS
)

from generator.document_styles import (
    apply_document_styles,
    format_title,
    format_heading,
)

from generator.utils import (
    format_chilean_date,
    has_real_content,
    is_not_applicable
)

from generator.section_classifier import (
    SectionClassifier
)

class DocumentGenerator:

    def _get_sections_for_modality(
        self,
        modality
    ):

        return MODALITY_SECTIONS.get(
            modality,
            ["identificacion", "informacion_adicional"]
        )

    def generate_document(
        self,
        df_group,
        group_info
    ):

        document = Document()

        apply_document_styles(document)

        self._add_header(document)

        self._add_cover_page(
            document,
            df_group,
            group_info
        )

        all_sections = SectionClassifier.build_section_map(
            df_group.columns
        )

        allowed_sections = self._get_sections_for_modality(
            group_info["modalidad"]
        )

        ordered_df = df_group.copy()

        grouped_records = {}

        for row in ordered_df.to_dict("records"):

            asesoria = row.get(
                COL_NOMBRE_ASESORIA,
                "Sin nombre"
            )

            grouped_records.setdefault(
                asesoria,
                []
            ).append(row)

        for asesoria_name, records in grouped_records.items():

            document.add_page_break()

            establishment_title = document.add_paragraph(
                f"Registro asociado a: {asesoria_name}"
            )

            format_heading(establishment_title)

            multiple_sessions = len(records) > 1

            for session_number, row in enumerate(
                records,
                start=1
            ):

                if multiple_sessions:

                    session_title = document.add_paragraph(
                        f"Sesión N.° {session_number}"
                    )

                    format_heading(session_title)

                date_text = format_chilean_date(
                    row.get(COL_FECHA)
                )

                document.add_paragraph(
                    f"Fecha de realización: {date_text}"
                )

                self._add_record(
                    document,
                    row,
                    all_sections,
                    allowed_sections
                )

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer

    def _add_header(self, document):

        section = document.sections[0]

        header = section.header

        paragraph = header.paragraphs[0]

        if LOGO_PATH.exists():

            run = paragraph.add_run()

            run.add_picture(
                str(LOGO_PATH),
                width=Inches(1.0)
            )

    def _add_cover_page(
        self,
        document,
        df_group,
        group_info
    ):

        title = document.add_paragraph(
            "Informe Individual Etapa De Implementación de la Asesoría"
        )

        format_title(title)

        table = document.add_table(
            rows=6,
            cols=2
        )

        establecimientos = df_group[
           COL_NOMBRE_ASESORIA
        ].nunique()

        table.style = "Table Grid"

        table.cell(0,0).text = "Región"
        table.cell(0,1).text = str(group_info["region"])

        table.cell(1,0).text = "DEPROV"
        table.cell(1,1).text = str(group_info["deprov"])

        table.cell(2,0).text = "Modalidad"
        table.cell(2,1).text = str(group_info["modalidad"])

        table.cell(3,0).text = "Asesor"
        table.cell(3,1).text = str(group_info["supervisor"])

        table.cell(4,0).text = "Total de asesorías"
        table.cell(4,1).text = str(len(df_group))

        table.cell(5,0).text = "Establecimientos asesorados"
        table.cell(5,1).text = str(establecimientos)

    def _add_record(
        self,
        document,
        row,
        all_sections,
        allowed_sections
    ):

        for section_name in allowed_sections:

            columns = all_sections.get(
                section_name,
                []
            )

            visible_columns = []

            for col in columns:

                value = row.get(col)

                if has_real_content(value):

                    if not is_not_applicable(value):

                        visible_columns.append(col)

            if not visible_columns:
                continue

            heading = document.add_paragraph(
                section_name.replace(
                    "_",
                    " "
                ).title()
            )

            format_heading(heading)

            table = document.add_table(
                rows=0,
                cols=2
            )

            table.style = "Table Grid"

            for col in visible_columns:

                value = row.get(col)

                r = table.add_row()

                r.cells[0].text = str(col)

                r.cells[1].text = str(value)
